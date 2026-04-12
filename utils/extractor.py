import os
import uuid
import PyPDF2
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
import openpyxl
import json

class ContentExtractor:
    """Extract content from various file types"""
    
    def extract_text_from_pdf(self, input_file):
        """Extract all text from PDF"""
        text = ""
        
        try:
            # Try PyMuPDF first (better text extraction)
            doc = fitz.open(input_file)
            
            for page in doc:
                text += page.get_text() + "\n\n"
            
            doc.close()
        
        except:
            # Fallback to PyPDF2
            with open(input_file, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n\n"
        
        return text.strip()
    
    def extract_images_from_pdf(self, input_file, output_folder):
        """Extract all images from PDF"""
        output_files = []
        
        try:
            doc = fitz.open(input_file)
            
            for page_num, page in enumerate(doc):
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # Save image
                    output_filename = f"{uuid.uuid4()}_page{page_num + 1}_img{img_index + 1}.{image_ext}"
                    output_path = os.path.join(output_folder, output_filename)
                    
                    with open(output_path, "wb") as img_file:
                        img_file.write(image_bytes)
                    
                    output_files.append(output_path)
            
            doc.close()
        
        except Exception as e:
            print(f"Error extracting images: {str(e)}")
        
        return output_files
    
    def extract_tables_from_pdf(self, input_file, output_folder):
        """Extract tables from PDF (basic implementation)"""
        try:
            import pdfplumber
            
            tables = []
            
            with pdfplumber.open(input_file) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_tables = page.extract_tables()
                    
                    for table_num, table in enumerate(page_tables):
                        # Save as CSV
                        output_filename = f"{uuid.uuid4()}_page{page_num + 1}_table{table_num + 1}.csv"
                        output_path = os.path.join(output_folder, output_filename)
                        
                        import csv
                        with open(output_path, 'w', newline='', encoding='utf-8') as f:
                            writer = csv.writer(f)
                            writer.writerows(table)
                        
                        tables.append(output_path)
            
            return tables
        
        except ImportError:
            print("pdfplumber not installed. Install with: pip install pdfplumber")
            return []
        except Exception as e:
            print(f"Error extracting tables: {str(e)}")
            return []
    
    def extract_metadata_from_pdf(self, input_file):
        """Extract PDF metadata"""
        metadata = {}
        
        try:
            with open(input_file, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if pdf_reader.metadata:
                    for key, value in pdf_reader.metadata.items():
                        # Remove leading '/' from keys
                        clean_key = key.lstrip('/')
                        metadata[clean_key] = value
                
                metadata['num_pages'] = len(pdf_reader.pages)
                metadata['encrypted'] = pdf_reader.is_encrypted
        
        except Exception as e:
            metadata['error'] = str(e)
        
        return metadata
    
    def extract_text_from_docx(self, input_file):
        """Extract text from Word document"""
        doc = Document(input_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text.strip()
    
    def extract_text_from_xlsx(self, input_file):
        """Extract text from Excel file"""
        wb = openpyxl.load_workbook(input_file)
        
        text = ""
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text += f"\n\n=== Sheet: {sheet_name} ===\n\n"
            
            for row in ws.iter_rows(values_only=True):
                row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                text += row_text + "\n"
        
        return text.strip()
    
    def extract_images_from_docx(self, input_file, output_folder):
        """Extract images from Word document"""
        output_files = []
        
        doc = Document(input_file)
        
        # Extract inline shapes
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_data = rel.target_part.blob
                
                # Determine file extension
                content_type = rel.target_part.content_type
                ext = content_type.split('/')[-1]
                if ext == 'jpeg':
                    ext = 'jpg'
                
                output_filename = f"{uuid.uuid4()}_docx_image.{ext}"
                output_path = os.path.join(output_folder, output_filename)
                
                with open(output_path, 'wb') as img_file:
                    img_file.write(img_data)
                
                output_files.append(output_path)
        
        return output_files
    
    def extract_hyperlinks_from_pdf(self, input_file):
        """Extract all hyperlinks from PDF"""
        links = []
        
        try:
            doc = fitz.open(input_file)
            
            for page_num, page in enumerate(doc):
                page_links = page.get_links()
                
                for link in page_links:
                    if 'uri' in link:
                        links.append({
                            'page': page_num + 1,
                            'url': link['uri']
                        })
            
            doc.close()
        
        except Exception as e:
            print(f"Error extracting links: {str(e)}")
        
        return links
    
    def extract_fonts_from_pdf(self, input_file):
        """Extract font information from PDF"""
        fonts = []
        
        try:
            doc = fitz.open(input_file)
            
            for page_num, page in enumerate(doc):
                font_list = page.get_fonts()
                
                for font in font_list:
                    fonts.append({
                        'page': page_num + 1,
                        'name': font[3],
                        'type': font[1]
                    })
            
            doc.close()
        
        except Exception as e:
            print(f"Error extracting fonts: {str(e)}")
        
        return fonts
    
    def extract_attachments_from_pdf(self, input_file, output_folder):
        """Extract file attachments from PDF"""
        output_files = []
        
        try:
            doc = fitz.open(input_file)
            
            # Get embedded files
            embedded_files = doc.embfile_names()
            
            for filename in embedded_files:
                file_info = doc.embfile_info(filename)
                file_content = doc.embfile_get(filename)
                
                output_path = os.path.join(output_folder, filename)
                
                with open(output_path, 'wb') as f:
                    f.write(file_content)
                
                output_files.append(output_path)
            
            doc.close()
        
        except Exception as e:
            print(f"Error extracting attachments: {str(e)}")
        
        return output_files
    
    def extract_form_data_from_pdf(self, input_file):
        """Extract form field data from PDF"""
        form_data = {}
        
        try:
            with open(input_file, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if '/AcroForm' in pdf_reader.trailer['/Root']:
                    fields = pdf_reader.get_fields()
                    
                    if fields:
                        for field_name, field_info in fields.items():
                            form_data[field_name] = field_info.get('/V', '')
        
        except Exception as e:
            form_data['error'] = str(e)
        
        return form_data
    
    def extract_page_as_image(self, input_file, output_folder, page_number=1, dpi=300):
        """Extract single PDF page as image"""
        output_filename = f"{uuid.uuid4()}_page{page_number}.png"
        output_path = os.path.join(output_folder, output_filename)
        
        try:
            doc = fitz.open(input_file)
            
            # Get page (0-indexed)
            page = doc[page_number - 1]
            
            # Render to image
            zoom = dpi / 72  # 72 is the default DPI
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            
            pix.save(output_path)
            doc.close()
        
        except Exception as e:
            print(f"Error extracting page: {str(e)}")
            return None
        
        return output_path
    
    def extract_all_pages_as_images(self, input_file, output_folder, dpi=300):
        """Extract all PDF pages as images"""
        output_files = []
        
        try:
            doc = fitz.open(input_file)
            
            for page_num in range(len(doc)):
                output_filename = f"{uuid.uuid4()}_page{page_num + 1}.png"
                output_path = os.path.join(output_folder, output_filename)
                
                page = doc[page_num]
                
                zoom = dpi / 72
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat)
                
                pix.save(output_path)
                output_files.append(output_path)
            
            doc.close()
        
        except Exception as e:
            print(f"Error extracting pages: {str(e)}")
        
        return output_files